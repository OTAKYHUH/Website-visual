# docx_to_html.py
"""
DOCX -> HTML (same folder), ready for VSCode "Run Python File".

- Preferred: uses `mammoth` for better DOCX->HTML.
  ✅ Images are embedded as base64 data URLs (NO assets folder).
- Fallback: uses `python-docx` basic converter (headings/paras/lists/tables),
  and can detect page breaks and insert <hr class="page-break">.

Run:
  - Put this file next to your .docx and press "Run Python File" in VSCode
  - Or: python docx_to_html.py "path/to/file.docx"
"""

from __future__ import annotations

import base64
import html
import os
import re
import sys
from pathlib import Path
from typing import Any, Iterator, Optional, Union, TYPE_CHECKING

from bs4 import BeautifulSoup  # pip install beautifulsoup4
from docx import Document      # pip install python-docx

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    Block = Union[Paragraph, Table]
else:
    Block = Any


def main() -> int:
    docx_path = get_input_path()
    if not docx_path.exists():
        print(f"[ERROR] DOCX not found: {docx_path}")
        return 2

    out_html_path = docx_path.with_suffix(".html")

    # If you really want to force the basic converter:
    # set environment variable FORCE_BASIC=1
    force_basic = os.getenv("FORCE_BASIC", "").strip() == "1"

    warnings: list[str] = []
    converter = "mammoth"

    if not force_basic:
        try:
            body_html, warnings = convert_with_mammoth_inline_images(docx_path)
        except Exception as e:
            print(f"[INFO] Mammoth not available/failed ({e}). Using basic converter.")
            body_html = convert_with_python_docx_basic(docx_path)
            warnings = ["Used basic converter (limited formatting)."]
            converter = "python-docx basic"
    else:
        body_html = convert_with_python_docx_basic(docx_path)
        warnings = ["FORCE_BASIC=1 enabled (limited formatting)."]
        converter = "python-docx basic"

    # Add heading IDs (anchors) for navigation
    body_html = add_heading_ids(body_html)

    final_html = wrap_html_document(
        body_html=body_html,
        title=docx_path.stem,
        converter=converter,
        warnings=warnings,
    )

    out_html_path.write_text(final_html, encoding="utf-8")
    print(f"[OK] Wrote: {out_html_path}")

    if warnings:
        print("[NOTES]")
        for w in warnings[:20]:
            print(f"  - {w}")

    return 0


def get_input_path() -> Path:
    # 1) Command-line path
    if len(sys.argv) >= 2:
        return Path(sys.argv[1]).expanduser().resolve()

    # 2) If script is next to a docx, pick it
    script_dir = Path(__file__).resolve().parent
    docxs = sorted(script_dir.glob("*.docx"))

    # Prefer your known name if present
    preferred = script_dir / "RP Booklet (Replanner Guideline).docx"
    if preferred.exists():
        return preferred

    if docxs:
        return docxs[0]

    # 3) Fallback default (will error cleanly)
    return preferred


# ----------------------------
# Preferred conversion: mammoth (INLINE images, no assets folder)
# ----------------------------
def convert_with_mammoth_inline_images(docx_path: Path) -> tuple[str, list[str]]:
    """
    Convert using mammoth. Images are embedded into HTML as base64 data URLs.
    No assets folder is created.
    """
    try:
        import mammoth  # pip install mammoth
    except ImportError as ie:
        raise RuntimeError("mammoth is not installed. Run: pip install mammoth") from ie

    def _inline_image(image) -> dict:
        content_type = (image.content_type or "image/png").lower()
        data = image.read()
        b64 = base64.b64encode(data).decode("ascii")
        return {"src": f"data:{content_type};base64,{b64}"}

    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(
            f,
            convert_image=mammoth.images.img_element(_inline_image),
        )

    warnings: list[str] = []
    for m in result.messages:
        try:
            warnings.append(f"{m.type}: {m.message}")
        except Exception:
            warnings.append(str(m))

    return result.value, warnings


# -----------------------------------
# Fallback conversion: python-docx basic
# -----------------------------------
def convert_with_python_docx_basic(docx_path: Path) -> str:
    """
    Basic DOCX->HTML:
    - Headings based on paragraph style "Heading X"
    - Paragraphs
    - Bullet/number lists (heuristic)
    - Tables
    - Attempts to detect page breaks and inserts <hr class="page-break">
    """
    doc = Document(str(docx_path))

    parts: list[str] = []
    list_mode: Optional[str] = None  # 'ul' or 'ol'

    def close_list():
        nonlocal list_mode
        if list_mode:
            parts.append(f"</{list_mode}>")
            list_mode = None

    for block in iter_block_items(doc):
        if is_table(block):
            close_list()
            parts.append(table_to_html(block))
            continue

        p = block  # paragraph
        text_html = paragraph_to_html(p)

        # Page breaks (best-effort)
        if paragraph_has_page_break(p):
            close_list()
            parts.append('<hr class="page-break" />')

        # Skip truly empty paragraphs (but keep page-break-only paragraphs)
        if not strip_html(text_html):
            continue

        heading_level = get_heading_level(p)
        if heading_level is not None:
            close_list()
            parts.append(f"<h{heading_level}>{text_html}</h{heading_level}>")
            continue

        li_kind = get_list_kind(p)  # 'ul'/'ol'/None
        if li_kind:
            if list_mode and list_mode != li_kind:
                close_list()
            if not list_mode:
                list_mode = li_kind
                parts.append(f"<{list_mode}>")
            parts.append(f"<li>{text_html}</li>")
        else:
            close_list()
            parts.append(f"<p>{text_html}</p>")

    close_list()
    return "\n".join(parts)


def paragraph_has_page_break(paragraph: Any) -> bool:
    """
    Best-effort detection of page breaks inside a paragraph.
    """
    try:
        r = paragraph._p
        # Detect explicit page breaks: <w:br w:type="page"/>
        brs = r.xpath(".//w:br[@w:type='page']")
        if brs:
            return True
        # Sometimes Word stores rendered page breaks
        last = r.xpath(".//w:lastRenderedPageBreak")
        return bool(last)
    except Exception:
        return False


# ----------------------------
# Post-process: add heading anchors
# ----------------------------
def add_heading_ids(body_html: str) -> str:
    soup = BeautifulSoup(body_html, "html.parser")
    used: set[str] = set()

    for tag in soup.find_all(re.compile(r"^h[1-6]$")):
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        slug = slugify(text)
        unique = make_unique_id(slug, used)
        tag["id"] = unique

    return str(soup)


def slugify(s: str) -> str:
    s = s.strip().lower()
    s = re.sub(r"[^\w\s-]", "", s)
    s = re.sub(r"[\s_-]+", "-", s)
    return s.strip("-") or "section"


def make_unique_id(base: str, used: set[str]) -> str:
    if base not in used:
        used.add(base)
        return base
    i = 2
    while f"{base}-{i}" in used:
        i += 1
    v = f"{base}-{i}"
    used.add(v)
    return v


# ----------------------------
# HTML wrapper
# ----------------------------
def wrap_html_document(body_html: str, title: str, converter: str, warnings: list[str]) -> str:
    warn_html = ""
    if warnings:
        warn_items = "\n".join(f"<li>{html.escape(w)}</li>" for w in warnings[:20])
        warn_html = f"""
        <details class="notes">
          <summary>Conversion notes</summary>
          <ul>{warn_items}</ul>
        </details>
        """

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{html.escape(title)}</title>
  <style>
    body {{
      font-family: Arial, sans-serif;
      line-height: 1.55;
      margin: 24px;
      max-width: 1000px;
    }}
    h1, h2, h3, h4, h5, h6 {{ margin-top: 1.2em; scroll-margin-top: 80px; }}
    table {{
      border-collapse: collapse;
      margin: 16px 0;
      width: 100%;
    }}
    th, td {{
      border: 1px solid #ccc;
      padding: 8px;
      vertical-align: top;
    }}
    img {{ max-width: 100%; height: auto; }}
    .meta {{
      color: #666;
      font-size: 12px;
      margin-bottom: 12px;
    }}
    .notes {{
      margin: 12px 0 18px;
      padding: 10px 12px;
      border: 1px solid #ddd;
      border-radius: 8px;
      background: #fafafa;
    }}
    hr.page-break {{
      border: none;
      border-top: 2px dashed #bbb;
      margin: 24px 0;
    }}
  </style>
</head>
<body>
  <div class="meta">Generated from DOCX using <b>{html.escape(converter)}</b>.</div>
  {warn_html}
  {body_html}
</body>
</html>
"""


# ----------------------------
# Basic converter helpers
# ----------------------------
def strip_html(s: str) -> str:
    return re.sub(r"<[^>]+>", "", s).strip()


def get_heading_level(paragraph: Any) -> Optional[int]:
    name = getattr(paragraph.style, "name", "") or ""
    m = re.match(r"Heading\s+(\d+)", name, flags=re.I)
    if m:
        lvl = int(m.group(1))
        return max(1, min(6, lvl))
    if name.lower() == "title":
        return 1
    return None


def get_list_kind(paragraph: Any) -> Optional[str]:
    name = (getattr(paragraph.style, "name", "") or "").lower()
    if "list bullet" in name or "bullet" in name:
        return "ul"
    if "list number" in name or "number" in name:
        return "ol"

    # Check XML numbering properties
    try:
        ppr = paragraph._p.pPr
        if ppr is not None and ppr.numPr is not None:
            return "ul"
    except Exception:
        pass

    return None


def paragraph_to_html(paragraph: Any) -> str:
    out = []
    for run in paragraph.runs:
        txt = html.escape(run.text)
        if not txt:
            continue
        if run.bold:
            txt = f"<strong>{txt}</strong>"
        if run.italic:
            txt = f"<em>{txt}</em>"
        if run.underline:
            txt = f"<u>{txt}</u>"
        out.append(txt)
    return "".join(out).strip()


def table_to_html(table: Any) -> str:
    rows_html = []
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            cell_parts = []
            for p in cell.paragraphs:
                ph = paragraph_to_html(p)
                if strip_html(ph):
                    cell_parts.append(ph)
            cell_html = "<br/>".join(cell_parts) if cell_parts else ""
            cells_html.append(f"<td>{cell_html}</td>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return "<table>\n" + "\n".join(rows_html) + "\n</table>"


def iter_block_items(parent: Any) -> Iterator[Block]:
    """
    Yield paragraphs and tables in document order.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def is_table(block: Any) -> bool:
    return hasattr(block, "rows") and hasattr(block, "columns")


if __name__ == "__main__":
    raise SystemExit(main())
